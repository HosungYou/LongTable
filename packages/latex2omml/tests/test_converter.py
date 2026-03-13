"""Tests for latex2omml converter."""

import pytest
from lxml import etree

from latex2omml import (
    add_display_equation,
    add_inline_equation,
    latex_to_omml_display,
    latex_to_omml_inline,
)

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _m(tag):
    return f"{{{MATH_NS}}}{tag}"


def _xml(elem):
    return etree.tostring(elem, pretty_print=True).decode()


# ---------------------------------------------------------------------------
# Basic expressions
# ---------------------------------------------------------------------------

class TestSimple:
    def test_single_letter(self):
        el = latex_to_omml_inline("x")
        runs = el.findall(f".//{_m('r')}")
        assert any(r.find(_m("t")).text == "x" for r in runs)

    def test_number(self):
        el = latex_to_omml_inline("42")
        runs = el.findall(f".//{_m('r')}")
        assert any(r.find(_m("t")).text == "42" for r in runs)

    def test_operator(self):
        el = latex_to_omml_inline("a + b")
        texts = [r.find(_m("t")).text for r in el.findall(f".//{_m('r')}")]
        assert any("+" in t for t in texts)

    def test_greek_letter(self):
        el = latex_to_omml_inline(r"\alpha")
        texts = [r.find(_m("t")).text for r in el.findall(f".//{_m('r')}")]
        assert any("\u03B1" in t for t in texts)

    def test_symbol(self):
        el = latex_to_omml_inline(r"\times")
        texts = [r.find(_m("t")).text for r in el.findall(f".//{_m('r')}")]
        assert any("\u00D7" in t for t in texts)


# ---------------------------------------------------------------------------
# Subscripts and superscripts
# ---------------------------------------------------------------------------

class TestScripts:
    def test_subscript(self):
        el = latex_to_omml_inline("x_i")
        assert el.find(f".//{_m('sSub')}") is not None

    def test_superscript(self):
        el = latex_to_omml_inline("x^2")
        assert el.find(f".//{_m('sSup')}") is not None

    def test_sub_and_sup(self):
        el = latex_to_omml_inline("x_{i}^{2}")
        assert el.find(f".//{_m('sSubSup')}") is not None

    def test_braced_subscript(self):
        el = latex_to_omml_inline("R_{b}")
        ssub = el.find(f".//{_m('sSub')}")
        assert ssub is not None
        sub = ssub.find(_m("sub"))
        texts = [r.find(_m("t")).text for r in sub.findall(f".//{_m('r')}")]
        assert "b" in texts


# ---------------------------------------------------------------------------
# Fractions
# ---------------------------------------------------------------------------

class TestFraction:
    def test_simple_frac(self):
        el = latex_to_omml_inline(r"\frac{a}{b}")
        frac = el.find(f".//{_m('f')}")
        assert frac is not None
        assert frac.find(_m("num")) is not None
        assert frac.find(_m("den")) is not None

    def test_nested_frac(self):
        el = latex_to_omml_inline(r"\frac{\frac{a}{b}}{c}")
        fracs = el.findall(f".//{_m('f')}")
        assert len(fracs) == 2

    def test_frac_with_text(self):
        el = latex_to_omml_inline(r"\frac{\text{numerator}}{\text{denominator}}")
        frac = el.find(f".//{_m('f')}")
        num_texts = [r.find(_m("t")).text for r in frac.find(_m("num")).findall(f".//{_m('r')}")]
        assert "numerator" in num_texts


# ---------------------------------------------------------------------------
# Text mode
# ---------------------------------------------------------------------------

class TestText:
    def test_text_plain_style(self):
        el = latex_to_omml_inline(r"\text{hello world}")
        runs = el.findall(f".//{_m('r')}")
        for r in runs:
            sty = r.find(f".//{_m('sty')}")
            if sty is not None and r.find(_m("t")).text == "hello world":
                assert sty.get(_m("val")) == "p"

    def test_text_with_underscore(self):
        el = latex_to_omml_inline(r"\text{adaptive\_offer}")
        texts = [r.find(_m("t")).text for r in el.findall(f".//{_m('r')}")]
        assert any("adaptive_offer" in t for t in texts)

    def test_text_with_subscript(self):
        el = latex_to_omml_inline(r"\text{Gap}_{\text{new}}")
        ssub = el.find(f".//{_m('sSub')}")
        assert ssub is not None


# ---------------------------------------------------------------------------
# Accents
# ---------------------------------------------------------------------------

class TestAccent:
    def test_hat(self):
        el = latex_to_omml_inline(r"\hat{p}")
        acc = el.find(f".//{_m('acc')}")
        assert acc is not None
        chr_elem = acc.find(f".//{_m('chr')}")
        assert chr_elem.get(_m("val")) == "\u0302"

    def test_bar(self):
        el = latex_to_omml_inline(r"\bar{x}")
        acc = el.find(f".//{_m('acc')}")
        assert acc is not None

    def test_hat_with_subscript(self):
        el = latex_to_omml_inline(r"\hat{p}_{ik}")
        ssub = el.find(f".//{_m('sSub')}")
        assert ssub is not None
        acc = ssub.find(f".//{_m('acc')}")
        assert acc is not None


# ---------------------------------------------------------------------------
# N-ary operators
# ---------------------------------------------------------------------------

class TestNary:
    def test_sum_with_limits(self):
        el = latex_to_omml_inline(r"\sum_{i=1}^{N} x_i")
        nary = el.find(f".//{_m('nary')}")
        assert nary is not None
        assert nary.find(_m("sub")) is not None
        assert nary.find(_m("sup")) is not None
        assert nary.find(_m("e")) is not None

    def test_sum_symbol(self):
        el = latex_to_omml_inline(r"\sum_{i=1}^{N} x")
        nary = el.find(f".//{_m('nary')}")
        chr_elem = nary.find(f".//{_m('chr')}")
        assert chr_elem.get(_m("val")) == "\u2211"

    def test_prod(self):
        el = latex_to_omml_inline(r"\prod_{k=1}^{K}")
        nary = el.find(f".//{_m('nary')}")
        chr_elem = nary.find(f".//{_m('chr')}")
        assert chr_elem.get(_m("val")) == "\u220F"


# ---------------------------------------------------------------------------
# Functions
# ---------------------------------------------------------------------------

class TestFunctions:
    def test_log_with_parens(self):
        el = latex_to_omml_inline(r"\log(x)")
        func = el.find(f".//{_m('func')}")
        assert func is not None
        d = func.find(f".//{_m('d')}")
        assert d is not None

    def test_function_name_plain(self):
        el = latex_to_omml_inline(r"\sin(x)")
        func = el.find(f".//{_m('func')}")
        fname = func.find(_m("fName"))
        sty = fname.find(f".//{_m('sty')}")
        assert sty.get(_m("val")) == "p"


# ---------------------------------------------------------------------------
# Square roots
# ---------------------------------------------------------------------------

class TestSqrt:
    def test_simple_sqrt(self):
        el = latex_to_omml_inline(r"\sqrt{x}")
        rad = el.find(f".//{_m('rad')}")
        assert rad is not None

    def test_nth_root(self):
        el = latex_to_omml_inline(r"\sqrt[3]{x}")
        rad = el.find(f".//{_m('rad')}")
        deg = rad.find(_m("deg"))
        texts = [r.find(_m("t")).text for r in deg.findall(f".//{_m('r')}")]
        assert "3" in texts


# ---------------------------------------------------------------------------
# Delimiters
# ---------------------------------------------------------------------------

class TestDelimiters:
    def test_parentheses(self):
        el = latex_to_omml_inline("(a + b)")
        d = el.find(f".//{_m('d')}")
        assert d is not None
        beg = d.find(f".//{_m('begChr')}")
        assert beg.get(_m("val")) == "("

    def test_brackets(self):
        el = latex_to_omml_inline("[a, b]")
        d = el.find(f".//{_m('d')}")
        assert d is not None
        beg = d.find(f".//{_m('begChr')}")
        assert beg.get(_m("val")) == "["


# ---------------------------------------------------------------------------
# Display vs inline
# ---------------------------------------------------------------------------

class TestModes:
    def test_display_has_oMathPara(self):
        el = latex_to_omml_display("x + y")
        assert el.tag == _m("oMathPara")
        assert el.find(_m("oMath")) is not None

    def test_inline_has_oMath(self):
        el = latex_to_omml_inline("x + y")
        assert el.tag == _m("oMath")


# ---------------------------------------------------------------------------
# Complex real-world equations (from Paper 4)
# ---------------------------------------------------------------------------

class TestPaper4Equations:
    def test_reliance_equation(self):
        latex = (
            r"R_b(\tau) = \frac{\text{Number of adaptive\_offer episodes "
            r"in window } \tau}{\text{Total number of episodes in window } \tau}"
        )
        el = latex_to_omml_display(latex)
        assert el.find(f".//{_m('f')}") is not None
        assert el.find(f".//{_m('sSub')}") is not None

    def test_gap_equation(self):
        el = latex_to_omml_display(r"\text{Gap}(\tau) = R_b(\tau) - P(\tau)")
        texts = [r.find(_m("t")).text for r in el.findall(f".//{_m('r')}")]
        assert any("Gap" in t for t in texts)

    def test_entropy_equation(self):
        latex = (
            r"E = 1 - \frac{-\sum_{i=1}^{N}\sum_{k=1}^{K} "
            r"\hat{p}_{ik} \log(\hat{p}_{ik})}{N \log(K)}"
        )
        el = latex_to_omml_display(latex)
        nary_elems = el.findall(f".//{_m('nary')}")
        assert len(nary_elems) >= 2
        assert el.find(f".//{_m('f')}") is not None
        assert el.find(f".//{_m('acc')}") is not None

    def test_ai_benefit_equation(self):
        latex = (
            r"\text{AI\_benefit}(\tau) = P_{\text{adaptive}}(\tau) "
            r"- P_{\text{non-adaptive}}(\tau)"
        )
        el = latex_to_omml_display(latex)
        ssubs = el.findall(f".//{_m('sSub')}")
        assert len(ssubs) >= 2


# ---------------------------------------------------------------------------
# python-docx integration
# ---------------------------------------------------------------------------

class TestDocxIntegration:
    def test_add_display_equation(self):
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        add_display_equation(p, r"\frac{a}{b}")
        assert p._element.find(f".//{_m('oMathPara')}") is not None

    def test_add_inline_equation(self):
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("The equation ")
        add_inline_equation(p, "E = mc^{2}")
        p.add_run(" is famous.")
        assert p._element.find(f".//{_m('oMath')}") is not None

    def test_save_and_verify(self, tmp_path):
        from docx import Document
        doc = Document()
        p = doc.add_paragraph()
        add_display_equation(p, r"\sum_{i=1}^{N} x_i^2")
        out = tmp_path / "test.docx"
        doc.save(str(out))
        assert out.exists()
        assert out.stat().st_size > 0
