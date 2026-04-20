from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


BASE_DIR = Path("/Users/hosung/long-table-refactoring-sync/docs/bjet-framework")
DOCX_PATH = BASE_DIR / "LongTable_BJET_Theoretical_Framework_APA7.docx"
FIG_PATH = Path("/Users/hosung/long-table-refactoring-sync/docs/assets/LongTable_BJET_Framework_Figure1.png")


def get_font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
        "/System/Library/Fonts/Supplemental/Georgia Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Georgia.ttf",
        "/System/Library/Fonts/Supplemental/Arial Bold.ttf" if bold else "/System/Library/Fonts/Supplemental/Arial.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def draw_centered_text(draw, box, text, font, fill):
    left, top, right, bottom = box
    width = right - left
    height = bottom - top
    words = text.split()
    lines = []
    current = ""
    for word in words:
        tentative = f"{current} {word}".strip()
        bbox = draw.textbbox((0, 0), tentative, font=font)
        if bbox[2] - bbox[0] <= width - 30:
            current = tentative
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)

    line_heights = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        line_heights.append(bbox[3] - bbox[1])
    total_height = sum(line_heights) + max(0, len(lines) - 1) * 8
    y = top + (height - total_height) / 2
    for i, line in enumerate(lines):
        bbox = draw.textbbox((0, 0), line, font=font)
        line_width = bbox[2] - bbox[0]
        x = left + (width - line_width) / 2
        draw.text((x, y), line, font=font, fill=fill)
        y += line_heights[i] + 8


def draw_arrow(draw, start, end, fill, width=6):
    draw.line([start, end], fill=fill, width=width)
    ex, ey = end
    sx, sy = start
    dx = ex - sx
    dy = ey - sy
    length = max((dx**2 + dy**2) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    arrow_size = 16
    p1 = (ex, ey)
    p2 = (
        ex - ux * arrow_size + px * (arrow_size * 0.6),
        ey - uy * arrow_size + py * (arrow_size * 0.6),
    )
    p3 = (
        ex - ux * arrow_size - px * (arrow_size * 0.6),
        ey - uy * arrow_size - py * (arrow_size * 0.6),
    )
    draw.polygon([p1, p2, p3], fill=fill)


def build_figure():
    img = Image.new("RGB", (1800, 1100), "white")
    draw = ImageDraw.Draw(img)
    title_font = get_font(44, bold=True)
    box_title_font = get_font(34, bold=True)
    box_body_font = get_font(28, bold=False)

    draw.text((90, 60), "Conceptual framework for LongTable in BJET", font=title_font, fill="#111111")

    colors = {
        "blue": "#DCEBFA",
        "green": "#DFF4E4",
        "gold": "#FFF0CC",
        "rose": "#FCE1E7",
        "line": "#4A4A4A",
    }

    boxes = {
        "human": (90, 220, 510, 520),
        "dialogic": (650, 170, 1160, 470),
        "agency": (650, 560, 1160, 860),
        "outcomes": (1310, 330, 1710, 700),
    }

    for name, box, fill in [
        ("human", boxes["human"], colors["blue"]),
        ("dialogic", boxes["dialogic"], colors["green"]),
        ("agency", boxes["agency"], colors["gold"]),
        ("outcomes", boxes["outcomes"], colors["rose"]),
    ]:
        draw.rounded_rectangle(box, radius=28, fill=fill, outline=colors["line"], width=4)

    draw_centered_text(draw, (120, 245, 480, 305), "Human mediation", box_title_font, "#111111")
    draw_centered_text(
        draw,
        (125, 315, 475, 500),
        "Visible curation\nProvenance and grounds\nContestability and revision",
        box_body_font,
        "#222222",
    )

    draw_centered_text(draw, (690, 195, 1120, 255), "Dialogic divergence", box_title_font, "#111111")
    draw_centered_text(
        draw,
        (690, 270, 1120, 445),
        "Multiple perspectives,\nframes, values, and reasons\nare juxtaposed rather than flattened",
        box_body_font,
        "#222222",
    )

    draw_centered_text(draw, (720, 585, 1090, 645), "Epistemic agency", box_title_font, "#111111")
    draw_centered_text(
        draw,
        (690, 660, 1120, 835),
        "Learners compare,\nchallenge, integrate,\nand justify claims",
        box_body_font,
        "#222222",
    )

    draw_centered_text(draw, (1340, 355, 1680, 415), "Educational outcomes", box_title_font, "#111111")
    draw_centered_text(
        draw,
        (1340, 430, 1680, 675),
        "Deeper reflection\nStronger justification quality\nLess unreflective AI reliance",
        box_body_font,
        "#222222",
    )

    draw_arrow(draw, (510, 350), (650, 280), colors["line"])
    draw_arrow(draw, (510, 390), (650, 690), colors["line"])
    draw_arrow(draw, (1160, 320), (1310, 430), colors["line"])
    draw_arrow(draw, (1160, 720), (1310, 600), colors["line"])

    note_font = get_font(24)
    draw.text(
        (90, 985),
        "Human mediation is the design condition that enables dialogic divergence and, in turn, supports epistemic agency.",
        font=note_font,
        fill="#333333",
    )

    img.save(FIG_PATH)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def set_default_style(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 2
    style.paragraph_format.space_after = Pt(0)

    for style_name in ["Title", "Subtitle", "Heading 1", "Heading 2"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_page_number(paragraph)


def add_title_page(doc):
    for _ in range(7):
        doc.add_paragraph("")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LongTable as a Human-Mediated Epistemic Tool:\nA BJET-Oriented Theoretical Framework")
    run.bold = True

    for _ in range(2):
        doc.add_paragraph("")

    for line in ["[Author Name]", "[Affiliation]", "[Author Note if Needed]"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line)

    doc.add_page_break()


def add_centered_heading(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    return p


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.5)
    p.add_run(text)
    return p


def add_references(doc, references):
    add_centered_heading(doc, "References")
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.add_run(ref)


def build_doc():
    doc = Document()
    set_default_style(doc)
    add_title_page(doc)

    add_centered_heading(doc, "Abstract")
    abstract = (
        "This concept note reframes LongTable for the British Journal of Educational Technology as a "
        "human-mediated epistemic tool rather than a human-like AI assistant. The framework argues that "
        "the educational value of generative AI arises when human mediation makes plurality, dialogue, "
        "and accountable judgment visible. To avoid an overloaded and incoherent theoretical base, the "
        "model anchors the paper in three linked constructs: dialogic pedagogy, epistemic agency, and "
        "human mediation informed by value-sensitive design. Together, these constructs define a narrower "
        "and more defensible account of how LongTable could support learning. The framework also clarifies "
        "why a BJET submission should foreground educational mechanisms, design principles, and learner "
        "judgment rather than anthropomorphism or system novelty alone."
    )
    p = doc.add_paragraph()
    p.add_run(abstract)
    p.paragraph_format.first_line_indent = Inches(0.5)

    keywords = doc.add_paragraph()
    keywords.add_run("Keywords: ").italic = True
    keywords.add_run("dialogic pedagogy, epistemic agency, human mediation, generative AI, educational technology")

    add_centered_heading(doc, "Theoretical Framework")
    add_body_paragraph(
        doc,
        "For a BJET submission, LongTable should be positioned as a human-mediated epistemic tool rather than "
        "as an AI system that merely appears human. The central claim is that generative AI becomes educationally "
        "valuable when human mediation structures multiple perspectives, dialogic reflection, and accountable "
        "judgment, not when automation simply increases speed or convenience (Turvey & Pachler, 2025; "
        "Wegerif & Casebourne, 2025). This move gives the paper a stable conceptual center and removes the "
        "need to combine too many loosely related theories at once."
    )
    add_body_paragraph(
        doc,
        "A parsimonious model for LongTable can therefore be organized around three levels. First, dialogic "
        "pedagogy explains learning as participation in responsive dialogue among perspectives rather than the "
        "passive reception of fixed answers (Wegerif & Casebourne, 2025). Second, epistemic agency captures "
        "the learner's active work of comparing, contesting, integrating, and revising claims in relation to "
        "AI-generated alternatives. In educational AI environments, this agency is strengthened when learners "
        "receive support for reflection and self-regulation rather than only fluent output (Xu et al., 2025). "
        "Third, human mediation and value-sensitive design clarify why educational systems should expose curation, "
        "provenance, and contestability instead of presenting outputs as if they were fully autonomous or neutral "
        "(Prieto et al., 2025)."
    )

    figure_num = doc.add_paragraph()
    figure_num.add_run("Figure 1").bold = True
    figure_title = doc.add_paragraph()
    figure_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = figure_title.add_run("Conceptual Framework for LongTable in BJET")
    title_run.italic = True

    fig = doc.add_paragraph()
    fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fig.add_run().add_picture(str(FIG_PATH), width=Inches(6.2))

    note = doc.add_paragraph()
    note.add_run("Note. ").italic = True
    note.add_run(
        "The figure summarizes the proposed causal logic for a BJET-oriented paper: visible human mediation "
        "supports dialogic divergence, which in turn supports learners' epistemic agency and educational outcomes."
    )

    add_body_paragraph(
        doc,
        "This framing leads to clearer BJET-facing research questions. A paper built from this framework can ask "
        "how visible human mediation shapes learners' epistemic agency, how students interpret curated traces of "
        "multiple perspectives, and whether this design reduces unreflective reliance while improving the quality "
        "of justification. Methodologically, the strongest path is a staged design-based program: qualitative "
        "elicitation with teachers and learners, iterative prototyping, and mixed-method evaluation combining "
        "interaction logs, artifacts, interviews, and a small number of validated measures (Prieto et al., 2025; "
        "Xu et al., 2025). This keeps the contribution educational and theoretical rather than merely technical."
    )
    add_body_paragraph(
        doc,
        "The main implication is that the BJET version of the paper should not try to prove everything at once. "
        "It should avoid building a single study around agency, trust, engagement, anthropomorphism, diversity, "
        "and learning outcomes simultaneously. Instead, the paper should defend one coherent proposition: "
        "human-mediated dialogic divergence is an educational design strategy that can strengthen learners' "
        "epistemic agency in generative AI environments."
    )

    references = [
        "Prieto, L., Viberg, O., Yip, J., & Topali, P. (2025). Aligning human values and educational technologies "
        "with value-sensitive design. British Journal of Educational Technology, 56(4), 1299-1310. "
        "https://doi.org/10.1111/bjet.13602",
        "Turvey, K., & Pachler, N. (2025). A topological exploration of convergence/divergence of human-mediated "
        "and algorithmically mediated pedagogy. British Journal of Educational Technology. Advance online publication. "
        "https://doi.org/10.1111/bjet.70007",
        "Wegerif, R., & Casebourne, I. (2025). A dialogic theoretical foundation for integrating generative AI into "
        "pedagogical design. British Journal of Educational Technology. Advance online publication. "
        "https://doi.org/10.1111/bjet.70026",
        "Xu, X., Qiao, L., Cheng, N., Liu, H., & Zhao, W. (2025). Enhancing self-regulated learning and learning "
        "experience in generative AI environments: The critical role of metacognitive support. British Journal of "
        "Educational Technology, 56(5), 1842-1863. https://doi.org/10.1111/bjet.13599",
    ]
    add_references(doc, references)

    doc.save(DOCX_PATH)


if __name__ == "__main__":
    build_figure()
    build_doc()
